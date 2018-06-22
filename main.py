# /usr/bin/env python
# coding=utf-8
from lib.MysqlClass import MysqlClass
from docx import Document
from docx.shared import Inches

if __name__ == '__main__':

    # 连接到mysql 数据库，提取到所有的表格
    db = MysqlClass("ip", "username", "password", "dbname", 3306)
    tables = db.getTables()

    # 文档样式设置
    document = Document()
    document.styles['Normal'].font.name = u'宋体'

    # 构建表格
    for table in tables:
        print("正在处理表:%s" % (table[0]))
        document.add_heading("%s [%s] " % (table[0].decode('utf8'), table[1].decode('utf8')), 3)

        # 读取列信息
        columns = db.getColumn(table[0])
        table = document.add_table(1, cols=3)
        table.style = 'Table Grid'

        # 增加标题标题
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u'名称'
        hdr_cells[1].text = u'类型'
        hdr_cells[2].text = u'备注'
        hdr_cells[2].width = Inches(4.25)

        # 增加表格数据
        for column in columns:
            row_cells = table.add_row().cells
            row_cells[0].text = column[0].decode('utf8')
            row_cells[1].text = column[1].decode('utf8')
            row_cells[2].text = column[2].decode('utf8')

        # 分页符
        document.add_page_break()

    # 保存文件
    try:
        document.save('db_document.docx')
    except:
        print('保存文件过程中发生了错误，请确认文件是否被占用')
    else:
        print('执行完毕！')
    finally:
        # 释放资源
        del db


